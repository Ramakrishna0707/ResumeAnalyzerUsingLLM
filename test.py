import streamlit as st
import PyPDF2
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_openai import OpenAI
from dotenv import load_dotenv
import openai
import os
import docx
from docx import Document
from io import BytesIO

def load_api_key():
    load_dotenv()
    openai.api_key = os.getenv('OPEN_API_KEY')

def extract_text_from_pdf(file):
    resume_text = ''
    pdf_reader = PyPDF2.PdfReader(file)
    for page in pdf_reader.pages:
        resume_text += page.extract_text()
    return resume_text

def extract_text_from_doc(file):
    resume_text = ''
    doc = docx.Document(file)
    for paragraph in doc.paragraphs:
        resume_text += paragraph.text + '\n'
    return resume_text

def extract_text_from_file(file):
    file_extension = file.name.split('.')[-1].lower()

    if file_extension == 'pdf':
        return extract_text_from_pdf(file)
    elif file_extension == 'docx':
        return extract_text_from_doc(file)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or Word document.")

def define_prompts():
    resume_skill_matching_prompt = PromptTemplate(
        input_variables=["resume_text", "job_description", "job_title"],
        template="""You are an experienced Technical Human Resource Manager. Your task is to review the provided resume against the given job description and job title Please share your professional evaluation on whether the candidate's profile aligns with the role. Highlight the strengths and weaknesses of the applicant in relation to the specified job requirements. {resume_text} {job_description} {job_title}"""
    )

    resume_percentage_prompt = PromptTemplate(
        input_variables=['resume_text', 'job_description'],
        template="""
        You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
        Your task is to evaluate the resume against the provided job description {job_description}. Give me the percentage of match if the resume matches
        the job description. First, the output should come as a percentage, then keywords missing, and last, final thoughts. {resume_text}
        """
    )

    resume_modification_prompt = PromptTemplate(
        input_variables= ['resume_text','job_description'],
        template = """you are an experienced resume writer. Your task is to modify the provided resume {resume_text} to better align with the given job description
        {job_description}. Incorporate relevant keywords and highlight the candidate's skills and experiences that matches the job requirements. Maintain the 
        overall structure and formatting of the resume for freshers role."""
    )
    return resume_skill_matching_prompt, resume_percentage_prompt,resume_modification_prompt

def create_llm_chains(api_key):
    llm = OpenAI(openai_api_key=api_key,max_tokens=2048)
    resume_skill_matching_prompt, resume_percentage_prompt, resume_modification_prompt = define_prompts()
    skill_matching_chain = LLMChain(llm=llm, prompt=resume_skill_matching_prompt, verbose=True)
    skill_percentage_chain = LLMChain(llm=llm, prompt=resume_percentage_prompt, verbose=True)
    resume_modification_chain = LLMChain(llm=llm, prompt=resume_modification_prompt, verbose=True)
    return skill_matching_chain, skill_percentage_chain, resume_modification_chain

def main():
    load_api_key()

    st.title("Resume Analyzer")
    st.text("Beat the ATS")
    job_description = st.text_area("Enter the Job Description")
    job_title = st.text_input("Enter the Job Title")
    resume_upload = st.file_uploader("Upload Resume: ", type=['PDF', 'DOCX'])

    # Render buttons
    strengths_weakness_button = st.button("Strengths and Weakness")
    percentage_match_button = st.button("Percentange matching with the job description")
    modify_resume_button = st.button("Modify Resume")
    download_resume_button = st.button("Download Modified Resume")

    if strengths_weakness_button:
        if job_title and job_description and resume_upload:
            try:
                resume_text = extract_text_from_file(resume_upload)
                skill_matching_chain, _, _ = create_llm_chains(openai.api_key)
                skill_matching_feedback = skill_matching_chain.run(resume_text=resume_text, job_description=job_description, job_title=job_title)

                with st.expander("Skill Matching Feedback"):
                    st.markdown("**Skill Matching Feedback**")
                    st.markdown(skill_matching_feedback)
            except ValueError as e:
                st.error(str(e))

    elif percentage_match_button:
        if job_description and resume_upload:
            try:
                resume_text = extract_text_from_file(resume_upload)
                _, skill_percentage_chain, _ = create_llm_chains(openai.api_key)
                skill_percentage = skill_percentage_chain.run(resume_text=resume_text, job_description=job_description)

                with st.expander("Overall Percentage Match"):
                    st.markdown("**Overall Percentage Match**")
                    st.markdown(skill_percentage)
            except ValueError as e:
                st.error(str(e))
    
    elif modify_resume_button:
        if job_description and resume_upload:
            try:
                resume_text = extract_text_from_file(resume_upload)
                _, _, resume_modification_chain = create_llm_chains(openai.api_key)
                modified_resume_text = resume_modification_chain.run(resume_text=resume_text, job_description=job_description)

                with st.expander("Modified Resume"):
                    st.markdown("**Modified Resume**")
                    st.markdown(modified_resume_text)

            except ValueError as e:
                st.error(str(e))
    
    elif download_resume_button:
        if job_description and resume_upload:
            try:
                resume_text = extract_text_from_file(resume_upload)
                _, _, resume_modification_chain = create_llm_chains(openai.api_key)
                modified_resume_text = resume_modification_chain.run(resume_text=resume_text, job_description=job_description)

                # Create a word document with the modified resume text
                doc = Document()
                doc.add_paragraph(modified_resume_text)

                # Save the word document to a BytesIO object
                output = BytesIO()
                doc.save(output)
                output.seek(0)

                # Download the word document
                st.download_button(
                    label="Download Modified Resume",
                    data=output.getvalue(),
                    file_name='modified_resume.docx',
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except ValueError as e:
                st.error(str(e))

if __name__ == "__main__":
    main()

